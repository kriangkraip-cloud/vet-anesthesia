import os, io
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session, joinedload
from datetime import datetime
from ..database import get_db
from .. import models, auth

router = APIRouter(prefix="/api/export", tags=["export"])

EXPORTS_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))), "exports")
os.makedirs(EXPORTS_DIR, exist_ok=True)


def _load_record(db: Session, record_id: int) -> models.AnestheticRecord:
    record = (
        db.query(models.AnestheticRecord)
        .options(
            joinedload(models.AnestheticRecord.patient),
            joinedload(models.AnestheticRecord.drug_entries),
            joinedload(models.AnestheticRecord.monitoring_entries),
            joinedload(models.AnestheticRecord.fluid_entries),
            joinedload(models.AnestheticRecord.emergency_events),
            joinedload(models.AnestheticRecord.surgical_procedures),
        )
        .filter(models.AnestheticRecord.id == record_id)
        .first()
    )
    if not record:
        raise HTTPException(status_code=404, detail="Record not found")
    return record


def _to_local(dt):
    """Convert a naive UTC datetime (from SQLite) to local time.
    Pure date objects are returned unchanged (no time component to convert)."""
    if dt is None:
        return None
    if not hasattr(dt, "strftime"):
        return dt
    from datetime import datetime as _DTType, timezone, timedelta
    import time as _time
    # date-only objects (no hour/minute) don't need timezone conversion
    if not isinstance(dt, _DTType):
        return dt
    utc_dt = dt.replace(tzinfo=timezone.utc)
    # time.timezone is seconds WEST of UTC (negative for UTC+7 → -25200)
    local_offset = _time.altzone if _time.daylight else _time.timezone
    local_dt = utc_dt - timedelta(seconds=local_offset)
    return local_dt.replace(tzinfo=None)


def _fmt_dt(dt) -> str:
    if dt is None:
        return ""
    local = _to_local(dt)
    return local.strftime("%H:%M") if hasattr(local, "strftime") else str(dt)


def _fmt_date(d) -> str:
    if d is None:
        return ""
    local = _to_local(d)
    return local.strftime("%d/%m/%Y") if hasattr(local, "strftime") else str(d)


def _calc_total_fluid(record, fluid_entries, monitoring_entries) -> float:
    """Total IV fluid = (maintenance rate × surgical time) + bolus."""
    main_fluid = next((f for f in fluid_entries
                       if f.fluid_type and f.fluid_type != "Fluid Bolus"), None)
    bolus_entry = next((f for f in fluid_entries
                        if f.fluid_type == "Fluid Bolus"), None)
    bolus_ml = (bolus_entry.total_volume or 0) if bolus_entry else 0

    # Manual override stored on the main fluid entry
    if main_fluid and getattr(main_fluid, "is_manual_volume", False) and main_fluid.total_volume:
        return float(main_fluid.total_volume)

    # Maintenance rate: from fluid entry, or last non-null monitoring fluid_rate
    rate = float(main_fluid.rate) if main_fluid and main_fluid.rate else 0
    if not rate:
        for m in sorted(monitoring_entries, key=lambda x: x.time, reverse=True):
            fr = getattr(m, "fluid_rate", None)
            if fr:
                rate = float(fr)
                break

    # Surgical time in hours
    surg_hours = 0.0
    if record.surgery_start and record.surgery_end:
        surg_hours = max(0.0, (record.surgery_end - record.surgery_start).total_seconds() / 3600)

    return rate * surg_hours + bolus_ml


def _generate_monitoring_chart(monitoring_entries) -> "io.BytesIO | None":
    try:
        import matplotlib
        matplotlib.use("Agg")
        # Patch broken macOS CoreText font scanner (matplotlib 3.9 bug with macOS 14+)
        try:
            import matplotlib.font_manager as _fm
            if hasattr(_fm, "_get_macos_fonts"):
                _fm._get_macos_fonts = lambda: []
        except Exception:
            pass
        import matplotlib.pyplot as plt
        import matplotlib.dates as mdates
    except ImportError:
        return None

    if not monitoring_entries or len(monitoring_entries) < 2:
        return None

    times = [m.time for m in monitoring_entries]

    fig, ax = plt.subplots(figsize=(13, 4))
    ax2 = ax.twinx()  # แกน Y ขวาสำหรับ Gas%

    vitals = [
        ("HR (bpm)",  [m.hr for m in monitoring_entries],           "#e63946", "o"),
        ("RR (/min)", [m.rr for m in monitoring_entries],           "#457b9d", "s"),
        ("SBP (mmHg)",[m.systolic_bp for m in monitoring_entries],  "#2a9d8f", "^"),
        ("DBP (mmHg)",[m.diastolic_bp for m in monitoring_entries], "#f4a261", "v"),
        ("MAP (mmHg)",[m.map_bp for m in monitoring_entries],       "#0077b6", "D"),
    ]

    has_data = False
    lines_all = []

    for label, values, color, marker in vitals:
        paired = [(t, v) for t, v in zip(times, values) if v is not None]
        if paired:
            ts, vs = zip(*paired)
            ln, = ax.plot(ts, vs, color=color, marker=marker, markersize=4,
                          linewidth=1.5, label=label)
            lines_all.append(ln)
            has_data = True

    # Gas% บนแกนขวา
    gas_vals = [m.gas_percent for m in monitoring_entries]
    gas_paired = [(t, v) for t, v in zip(times, gas_vals) if v is not None]
    if gas_paired:
        ts_g, vs_g = zip(*gas_paired)
        ln_gas, = ax2.plot(ts_g, vs_g, color="#9b5de5", marker="p", markersize=5,
                           linewidth=2, linestyle="--", label="Gas% (%)")
        lines_all.append(ln_gas)
        has_data = True

    if not has_data:
        plt.close(fig)
        return None

    ax.xaxis.set_major_formatter(mdates.DateFormatter("%H:%M"))
    plt.xticks(rotation=30, fontsize=8)
    ax.set_xlabel("Time", fontsize=9)
    ax.set_ylabel("HR / RR / BP (mmHg)", fontsize=9)
    ax2.set_ylabel("Gas %", fontsize=9, color="#9b5de5")
    ax2.tick_params(axis="y", labelcolor="#9b5de5", labelsize=8)
    ax2.set_ylim(0, 8)
    ax.set_title("Intraoperative Monitoring  —  HR · RR · SBP · DBP · MAP · Gas%", fontsize=10, pad=6)
    labels_all = [l.get_label() for l in lines_all]
    ax.legend(lines_all, labels_all, loc="upper right", fontsize=8, ncol=6, framealpha=0.7)
    ax.grid(True, alpha=0.25)
    ax.set_ylim(bottom=0)
    plt.tight_layout()

    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=110, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf


@router.get("/{record_id}/pdf")
async def export_pdf(
    record_id: int,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(auth.get_current_user)
):
    record = _load_record(db, record_id)
    recorder_name = current_user.full_name or current_user.username
    buf = _build_pdf(record, recorder_name=recorder_name)
    filename = f"Anesthesia_{record.patient.hn}_{_fmt_date(record.record_date).replace('/', '-')}.pdf"

    log = models.ExportHistory(record_id=record_id, export_type="pdf", exported_by_id=current_user.id)
    db.add(log)
    if record.status == "completed":
        record.status = "exported"
    db.commit()

    return StreamingResponse(
        buf,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.get("/{record_id}/docx")
async def export_docx(
    record_id: int,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(auth.get_current_user)
):
    record = _load_record(db, record_id)
    recorder_name = current_user.full_name or current_user.username
    buf = _build_docx(record, recorder_name=recorder_name)
    filename = f"Anesthesia_{record.patient.hn}_{_fmt_date(record.record_date).replace('/', '-')}.docx"

    log = models.ExportHistory(record_id=record_id, export_type="docx", exported_by_id=current_user.id)
    db.add(log)
    if record.status == "completed":
        record.status = "exported"
    db.commit()

    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


def _build_pdf(record: models.AnestheticRecord, recorder_name: str = "") -> io.BytesIO:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.units import cm
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, HRFlowable, Image as RLImage
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, topMargin=1.5*cm, bottomMargin=1.5*cm, leftMargin=1.5*cm, rightMargin=1.5*cm)
    styles = getSampleStyleSheet()
    elements = []

    title_style = ParagraphStyle("title", parent=styles["Title"], fontSize=14, spaceAfter=4)
    h2 = ParagraphStyle("h2", parent=styles["Heading2"], fontSize=11, spaceAfter=2, spaceBefore=8)
    normal = styles["Normal"]
    small = ParagraphStyle("small", parent=normal, fontSize=8)

    TEAL = colors.HexColor("#0077b6")
    LIGHT = colors.HexColor("#e8f4f8")

    def section(title):
        elements.append(Spacer(1, 0.3*cm))
        elements.append(Paragraph(title, h2))
        elements.append(HRFlowable(width="100%", thickness=1, color=TEAL, spaceAfter=4))

    def kv_table(rows, cols=2):
        data = []
        row_data = []
        for label, value in rows:
            row_data.append(Paragraph(f"<b>{label}:</b>", small))
            row_data.append(Paragraph(str(value or ""), small))
            if len(row_data) == cols * 2:
                data.append(row_data)
                row_data = []
        if row_data:
            while len(row_data) < cols * 2:
                row_data.extend([Paragraph("", small), Paragraph("", small)])
            data.append(row_data)
        col_widths = [3*cm, 5.5*cm] * cols
        t = Table(data, colWidths=col_widths)
        t.setStyle(TableStyle([
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ("TOPPADDING", (0, 0), (-1, -1), 3),
        ]))
        elements.append(t)

    p = record.patient

    # Header
    elements.append(Paragraph("VETERINARY ANESTHETIC RECORD", title_style))
    elements.append(Paragraph(f"Generated: {datetime.now().strftime('%d/%m/%Y %H:%M')}", small))
    if recorder_name:
        elements.append(Paragraph(f"<b>Recorded by:</b> {recorder_name}", small))
    elements.append(Spacer(1, 0.3*cm))

    # Patient info
    section("Patient Information")
    kv_table([
        ("HN", p.hn), ("Patient Name", p.name),
        ("Species", p.species), ("Breed", p.breed),
        ("Sex", p.sex), ("Neuter Status", p.neuter_status),
        ("Age", p.age), ("Body Weight", f"{p.weight} kg" if p.weight else ""),
        ("Owner", p.owner_name), ("Date", _fmt_date(record.record_date)),
        ("Surgeon", record.surgeon), ("Anesthesiologist", record.anesthesiologist),
        ("Assistant", record.assistant), ("Diagnosis", record.diagnosis),
        ("Procedure", record.surgical_procedure), ("Status", record.status.upper()),
    ], cols=2)

    # Pre-anesthetic eval
    section("Pre-Anesthetic Evaluation")
    kv_table([
        ("ASA Status", record.asa_status), ("BW at Record", f"{record.weight_at_record} kg" if record.weight_at_record else ""),
        ("PCV (%)", record.pcv), ("BUN (mg/dL)", record.bun),
        ("Creatinine", record.creatinine), ("AST", record.ast),
        ("ALT", record.alt), ("CRT", record.crt),
        ("HR (bpm)", record.hr_pre), ("RR (/min)", record.rr_pre),
        ("Pulse Quality", record.pulse_quality), ("Temp (°C)", record.temp_pre),
    ], cols=2)

    if record.current_medications:
        elements.append(Paragraph(f"<b>Current Medications:</b> {record.current_medications}", small))
    if record.clinical_notes:
        elements.append(Paragraph(f"<b>Clinical Notes:</b> {record.clinical_notes}", small))
    if record.anesthetic_concerns:
        elements.append(Paragraph(f"<b>Anesthetic Concerns:</b> {record.anesthetic_concerns}", small))

    # Drug sections
    drug_types = [
        ("premedication", "Premedication"),
        ("induction", "Induction"),
        ("antibiotic", "Prophylactic Antibiotic"),
        ("local_anesthesia", "Local / Regional Anesthesia"),
    ]
    for dtype, dtitle in drug_types:
        drugs = [d for d in record.drug_entries if d.entry_type == dtype]
        if drugs:
            section(dtitle)
            headers = ["Drug", "Dose", "Conc.", "Volume", "Route", "Time", "Note"]
            data = [headers]
            for d in drugs:
                data.append([
                    d.drug_name or "",
                    f"{d.dose} {d.dose_unit}" if d.dose else "",
                    f"{d.concentration} mg/mL" if d.concentration else "",
                    f"{d.calculated_volume:.2f} mL" if d.calculated_volume else "",
                    d.route or "",
                    _fmt_dt(d.time),
                    d.note or "",
                ])
            t = Table(data, colWidths=[3*cm, 2.5*cm, 2.5*cm, 2.5*cm, 1.8*cm, 1.8*cm, 3.4*cm])
            t.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), TEAL),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTSIZE", (0, 0), (-1, -1), 7),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, LIGHT]),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("TOPPADDING", (0, 0), (-1, -1), 3),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ]))
            elements.append(t)

    # Airway & Equipment
    section("Airway & Equipment")
    kv_table([
        ("Premed Quality", record.premed_quality), ("Induction Quality", record.induction_quality),
        ("O2 Flow (L/min)", record.o2_flow_rate), ("Machine", record.anesthetic_machine),
        ("Breathing System", record.breathing_system), ("Gas/Vaporizer", record.vaporizer_gas),
        ("Airway Device", record.airway_device), ("ETT Size (mm)", record.ett_size),
    ], cols=2)
    if record.intubation_note:
        elements.append(Paragraph(f"<b>Intubation Note:</b> {record.intubation_note}", small))
    if record.induction_complication:
        elements.append(Paragraph(f"<b>Complication:</b> {record.induction_complication}", small))

    # Monitoring chart
    chart_buf = _generate_monitoring_chart(record.monitoring_entries)
    if chart_buf:
        section("Monitoring Chart (HR · RR · SBP · DBP · MAP · Gas%)")
        chart_img = RLImage(chart_buf, width=17*cm, height=5.5*cm)
        elements.append(chart_img)
        elements.append(Spacer(1, 0.2*cm))

    # Monitoring table
    if record.monitoring_entries:
        section("Intraoperative Monitoring Log")
        mon_headers = ["Time", "Gas%", "HR", "RR", "SpO2", "SBP", "DBP", "MAP", "EtCO2", "Temp", "O2", "Fluid", "Mode", "Note"]
        mon_data = [mon_headers]
        for m in record.monitoring_entries:
            fluid_rate_val = getattr(m, "fluid_rate", None)
            mon_data.append([
                _fmt_dt(m.time),
                str(m.gas_percent or ""), str(int(m.hr) if m.hr else ""),
                str(int(m.rr) if m.rr else ""), str(m.spo2 or ""),
                str(int(m.systolic_bp) if m.systolic_bp else ""),
                str(int(m.diastolic_bp) if m.diastolic_bp else ""),
                str(int(m.map_bp) if m.map_bp else ""),
                str(m.etco2 or ""), str(m.body_temp or ""),
                str(m.o2_flow or ""),
                str(int(fluid_rate_val) if fluid_rate_val else ""),
                (m.ventilation_mode or "")[:4],
                (m.note or "")[:18],
            ])
        col_w = [1.3*cm, 1*cm, 1*cm, 1*cm, 1*cm, 1*cm, 1*cm, 1*cm, 1*cm, 1*cm, 1*cm, 1.2*cm, 1.1*cm, 2.1*cm]
        t = Table(mon_data, colWidths=col_w)
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), TEAL),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTSIZE", (0, 0), (-1, -1), 7),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, LIGHT]),
            ("TOPPADDING", (0, 0), (-1, -1), 2),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
        ]))
        elements.append(t)

    # Fluid Therapy
    main_fluids = [f for f in record.fluid_entries if f.fluid_type and f.fluid_type != "Fluid Bolus"]
    total_fluid_ml = _calc_total_fluid(record, record.fluid_entries, record.monitoring_entries)
    if main_fluids or total_fluid_ml > 0:
        section("Fluid Therapy")
        fluid_type_str = main_fluids[0].fluid_type if main_fluids else "—"
        elements.append(Paragraph(f"<b>IV Fluid Type:</b> {fluid_type_str}", small))
        bolus_entry = next((f for f in record.fluid_entries if f.fluid_type == "Fluid Bolus"), None)
        bolus_ml = bolus_entry.total_volume or 0 if bolus_entry else 0
        main_rate = main_fluids[0].rate if main_fluids and main_fluids[0].rate else ""
        if main_rate:
            elements.append(Paragraph(f"<b>IV Rate:</b> {main_rate} mL/hr", small))
        if bolus_ml:
            elements.append(Paragraph(f"<b>Fluid Bolus:</b> {bolus_ml:.0f} mL", small))
        elements.append(Paragraph(f"<b>Total IV Volume:</b> {total_fluid_ml:.1f} mL", small))

    # CRI
    cri_entries = [d for d in record.drug_entries if d.entry_type == "cri"]
    if cri_entries:
        section("CRI Medications")
        c_headers = ["Time", "Drug", "Rate (mL/hr)", "Note"]
        c_data = [c_headers]
        for c in cri_entries:
            c_data.append([
                _fmt_dt(c.time),
                c.drug_name or "",
                f"{c.rate}" if c.rate is not None else "",
                c.note or "",
            ])
        t = Table(c_data, colWidths=[2.5*cm, 6*cm, 4*cm, 4*cm])
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), TEAL),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTSIZE", (0, 0), (-1, -1), 8),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("TOPPADDING", (0, 0), (-1, -1), 3),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
        ]))
        elements.append(t)

    # Emergency events
    if record.emergency_events:
        section("Emergency Drug & Event Log")
        e_headers = ["Time", "Event", "Drug", "Dose", "Volume", "Route", "Response", "Note"]
        e_data = [e_headers]
        for e in record.emergency_events:
            e_data.append([
                _fmt_dt(e.time), e.event_type or "", e.drug_name or "",
                f"{e.dose} {e.dose_unit}" if e.dose else "",
                f"{e.volume:.2f} mL" if e.volume else "",
                e.route or "", (e.response or "")[:30], (e.note or "")[:30],
            ])
        t = Table(e_data, colWidths=[1.5*cm, 2.5*cm, 2.5*cm, 2*cm, 1.8*cm, 1.5*cm, 2.5*cm, 2*cm])
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#e63946")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTSIZE", (0, 0), (-1, -1), 7),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("TOPPADDING", (0, 0), (-1, -1), 3),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
        ]))
        elements.append(t)

    # Surgical timing
    section("Timing")
    kv_table([
        ("Anesthesia Start", _fmt_dt(record.anesthesia_start)), ("Anesthesia End", _fmt_dt(record.anesthesia_end)),
        ("Procedure Start", _fmt_dt(record.surgery_start)), ("Procedure End", _fmt_dt(record.surgery_end)),
    ], cols=2)

    # Recovery
    section("Recovery")
    kv_table([
        ("Extubation Time", _fmt_dt(record.extubation_time)), ("Sternal Time", _fmt_dt(record.sternal_time)),
        ("Standing Time", _fmt_dt(record.standing_time)), ("Recovery Quality", record.recovery_quality),
    ], cols=2)
    if record.recovery_complications:
        elements.append(Paragraph(f"<b>Complications:</b> {record.recovery_complications}", small))
    if record.postop_pain_management:
        elements.append(Paragraph(f"<b>Postop Pain Mgmt:</b> {record.postop_pain_management}", small))
    if record.postop_medications:
        elements.append(Paragraph(f"<b>Postop Medications:</b> {record.postop_medications}", small))
    if record.final_note:
        elements.append(Paragraph(f"<b>Final Note:</b> {record.final_note}", small))

    # Signatures
    section("Signatures")
    sig_data = [
        ["Anesthesiologist:", record.anesthesiologist or "", "Surgeon:", record.surgeon or ""],
        ["", "_________________________", "", "_________________________"],
        ["Date:", _fmt_date(record.record_date), "Date:", _fmt_date(record.record_date)],
    ]
    if recorder_name:
        sig_data.append(["Recorded by:", recorder_name, "", ""])
    t = Table(sig_data, colWidths=[3*cm, 6*cm, 3*cm, 6*cm])
    t.setStyle(TableStyle([
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
    ]))
    elements.append(t)

    doc.build(elements)
    buf.seek(0)
    return buf


def _build_docx(record: models.AnestheticRecord, recorder_name: str = "") -> io.BytesIO:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    p = record.patient

    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    def add_heading(text, level=1):
        h = doc.add_heading(text, level=level)
        h.runs[0].font.color.rgb = RGBColor(0, 119, 182)

    def add_kv(label, value, bold_label=True):
        p_obj = doc.add_paragraph()
        run = p_obj.add_run(f"{label}: ")
        run.bold = bold_label
        p_obj.add_run(str(value or ""))
        p_obj.paragraph_format.space_after = Pt(2)

    def add_table(headers, rows):
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr[i].text = h
            hdr[i].paragraphs[0].runs[0].bold = True
            hdr[i].paragraphs[0].runs[0].font.size = Pt(8)
        for row_data in rows:
            row = table.add_row().cells
            for i, cell_val in enumerate(row_data):
                row[i].text = str(cell_val or "")
                row[i].paragraphs[0].runs[0].font.size = Pt(8)

    doc.add_heading("VETERINARY ANESTHETIC RECORD", 0)
    gen_para = doc.add_paragraph(f"Generated: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    gen_para.runs[0].font.size = Pt(9)
    if recorder_name:
        rec_para = doc.add_paragraph()
        run = rec_para.add_run("Recorded by: ")
        run.bold = True
        run.font.size = Pt(9)
        rec_para.add_run(recorder_name).font.size = Pt(9)

    add_heading("Patient Information", 1)
    add_kv("HN", p.hn)
    add_kv("Patient Name", p.name)
    add_kv("Species / Breed", f"{p.species or ''} / {p.breed or ''}")
    add_kv("Sex / Neuter", f"{p.sex or ''} / {p.neuter_status or ''}")
    add_kv("Age", p.age)
    add_kv("Body Weight", f"{p.weight} kg" if p.weight else "")
    add_kv("Owner", p.owner_name)
    add_kv("Date", _fmt_date(record.record_date))
    add_kv("Surgeon", record.surgeon)
    add_kv("Anesthesiologist", record.anesthesiologist)
    add_kv("Assistant", record.assistant)
    add_kv("Diagnosis", record.diagnosis)
    add_kv("Procedure", record.surgical_procedure)

    add_heading("Pre-Anesthetic Evaluation", 1)
    add_kv("ASA Status", record.asa_status)
    add_kv("PCV / BUN / Creatinine", f"{record.pcv or ''}% / {record.bun or ''} / {record.creatinine or ''}")
    add_kv("AST / ALT", f"{record.ast or ''} / {record.alt or ''}")
    add_kv("HR / RR", f"{record.hr_pre or ''} bpm / {record.rr_pre or ''} /min")
    add_kv("CRT / Pulse", f"{record.crt or ''} / {record.pulse_quality or ''}")
    add_kv("Temperature", f"{record.temp_pre or ''}°C")
    add_kv("Current Medications", record.current_medications)
    add_kv("Clinical Notes", record.clinical_notes)
    add_kv("Anesthetic Concerns", record.anesthetic_concerns)

    for dtype, dtitle in [("premedication", "Premedication"), ("induction", "Induction"),
                           ("antibiotic", "Prophylactic Antibiotic"), ("local_anesthesia", "Local Anesthesia")]:
        drugs = [d for d in record.drug_entries if d.entry_type == dtype]
        if drugs:
            add_heading(dtitle, 1)
            rows = [[d.drug_name, f"{d.dose} {d.dose_unit}" if d.dose else "",
                     f"{d.concentration} mg/mL" if d.concentration else "",
                     f"{d.calculated_volume:.2f} mL" if d.calculated_volume else "",
                     d.route, _fmt_dt(d.time), d.note] for d in drugs]
            add_table(["Drug", "Dose", "Conc.", "Volume", "Route", "Time", "Note"], rows)

    add_heading("Airway & Equipment", 1)
    add_kv("Premed Quality", record.premed_quality)
    add_kv("Induction Quality", record.induction_quality)
    add_kv("O2 Flow", f"{record.o2_flow_rate} L/min" if record.o2_flow_rate else "")
    add_kv("Machine / System", f"{record.anesthetic_machine or ''} / {record.breathing_system or ''}")
    add_kv("Gas / Vaporizer", record.vaporizer_gas)
    add_kv("Airway Device / ETT", f"{record.airway_device or ''} {record.ett_size or ''} mm")
    if record.intubation_note:
        add_kv("Intubation Note", record.intubation_note)

    # Monitoring chart
    chart_buf = _generate_monitoring_chart(record.monitoring_entries)
    if chart_buf:
        add_heading("Monitoring Chart (HR · RR · SBP · DBP · MAP · Gas%)", 1)
        try:
            doc.add_picture(chart_buf, width=Cm(16))
        except Exception:
            pass

    if record.monitoring_entries:
        add_heading("Intraoperative Monitoring Log", 1)
        rows = []
        for m in record.monitoring_entries:
            fluid_rate_val = getattr(m, "fluid_rate", None)
            rows.append([_fmt_dt(m.time), m.gas_percent, m.hr, m.rr, m.spo2,
                         m.systolic_bp, m.diastolic_bp, m.map_bp, m.etco2,
                         m.body_temp, m.o2_flow,
                         f"{fluid_rate_val}" if fluid_rate_val is not None else "",
                         m.ventilation_mode, m.note])
        add_table(["Time", "Gas%", "HR", "RR", "SpO2", "SBP", "DBP", "MAP", "EtCO2", "Temp", "O2", "Fluid", "Mode", "Note"], rows)

    # Fluid therapy (simplified)
    main_fluids = [f for f in record.fluid_entries if f.fluid_type and f.fluid_type != "Fluid Bolus"]
    total_fluid_ml = _calc_total_fluid(record, record.fluid_entries, record.monitoring_entries)
    if main_fluids or total_fluid_ml > 0:
        add_heading("Fluid Therapy", 1)
        if main_fluids:
            add_kv("IV Fluid Type", main_fluids[0].fluid_type)
            if main_fluids[0].rate:
                add_kv("IV Rate", f"{main_fluids[0].rate} mL/hr")
        bolus_entry = next((f for f in record.fluid_entries if f.fluid_type == "Fluid Bolus"), None)
        if bolus_entry and bolus_entry.total_volume:
            add_kv("Fluid Bolus", f"{bolus_entry.total_volume:.0f} mL")
        add_kv("Total IV Volume", f"{total_fluid_ml:.1f} mL")

    cri_entries = [d for d in record.drug_entries if d.entry_type == "cri"]
    if cri_entries:
        add_heading("CRI Medications", 1)
        rows = [[_fmt_dt(c.time), c.drug_name, f"{c.rate} mL/hr" if c.rate is not None else "", c.note]
                for c in cri_entries]
        add_table(["Time", "Drug", "Rate (mL/hr)", "Note"], rows)

    if record.emergency_events:
        add_heading("Emergency Drug & Event Log", 1)
        rows = [[_fmt_dt(e.time), e.event_type, e.drug_name,
                 f"{e.dose} {e.dose_unit}" if e.dose else "", e.volume,
                 e.route, e.response, e.note] for e in record.emergency_events]
        add_table(["Time", "Event", "Drug", "Dose", "Vol (mL)", "Route", "Response", "Note"], rows)

    add_heading("Timing", 1)
    add_kv("Anesthesia Start / End", f"{_fmt_dt(record.anesthesia_start)} / {_fmt_dt(record.anesthesia_end)}")
    add_kv("Procedure Start / End", f"{_fmt_dt(record.surgery_start)} / {_fmt_dt(record.surgery_end)}")

    add_heading("Recovery", 1)
    add_kv("Extubation Time", _fmt_dt(record.extubation_time))
    add_kv("Sternal Time", _fmt_dt(record.sternal_time))
    add_kv("Standing Time", _fmt_dt(record.standing_time))
    add_kv("Recovery Quality", record.recovery_quality)
    add_kv("Complications", record.recovery_complications)
    add_kv("Postop Pain Management", record.postop_pain_management)
    add_kv("Postop Medications", record.postop_medications)
    add_kv("Final Note", record.final_note)

    # Signatures
    add_heading("Signatures", 1)
    sig_rows = 3 + (1 if recorder_name else 0)
    sig_table = doc.add_table(rows=sig_rows, cols=4)
    sig_table.cell(0, 0).text = "Anesthesiologist:"
    sig_table.cell(0, 1).text = record.anesthesiologist or ""
    sig_table.cell(0, 2).text = "Surgeon:"
    sig_table.cell(0, 3).text = record.surgeon or ""
    sig_table.cell(1, 1).text = "_" * 25
    sig_table.cell(1, 3).text = "_" * 25
    sig_table.cell(2, 0).text = "Date:"
    sig_table.cell(2, 1).text = _fmt_date(record.record_date)
    sig_table.cell(2, 2).text = "Date:"
    sig_table.cell(2, 3).text = _fmt_date(record.record_date)
    if recorder_name:
        sig_table.cell(3, 0).text = "Recorded by:"
        sig_table.cell(3, 1).text = recorder_name

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
